"""
M4 Acceptance Tests: Extraction (Text/PDF/DOCX)

Tests use tmp_path only per PRACTICES.md section 4.3.
"""
import hashlib
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter

from src.assistant.ingestion.extractors import (
    ExtractionResult,
    extract_docx,
    extract_pdf,
    extract_text,
)


def test_extract_text_basic():
    """Test basic text extraction."""
    result = extract_text("Hello World", label="test")

    assert result.text == "Hello World"
    assert result.input_kind == "text"
    assert result.input_name == "test"
    assert result.input_hash == hashlib.sha256("Hello World".encode()).hexdigest()


def test_extract_text_strips_whitespace():
    """Test text extraction strips leading/trailing whitespace."""
    result = extract_text("  Hello World  \n\n")

    assert result.text == "Hello World"


def test_extract_text_empty_raises_error():
    """Test empty text raises ValueError."""
    with pytest.raises(ValueError, match="empty"):
        extract_text("")

    with pytest.raises(ValueError, match="empty"):
        extract_text("   \n\n  ")


def test_extract_text_deterministic_hash():
    """Test hash is deterministic for same content."""
    result1 = extract_text("Same content")
    result2 = extract_text("Same content")

    assert result1.input_hash == result2.input_hash


def test_extract_text_different_hash_for_different_content():
    """Test hash differs for different content."""
    result1 = extract_text("Content A")
    result2 = extract_text("Content B")

    assert result1.input_hash != result2.input_hash


def test_extract_pdf(tmp_path):
    """Test PDF extraction using tmp_path."""
    pdf_path = tmp_path / "test.pdf"

    writer = PdfWriter()
    from pypdf import PageObject
    from pypdf.generic import RectangleObject
    import io
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.pagesizes import letter

    # Create a simple PDF with reportlab
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    c.drawString(100, 700, "Test PDF content page 1")
    c.showPage()
    c.drawString(100, 700, "Test PDF content page 2")
    c.showPage()
    c.save()

    pdf_path.write_bytes(buf.getvalue())

    result = extract_pdf(pdf_path)

    assert result.input_kind == "pdf"
    assert result.input_name == "test.pdf"
    assert "Test PDF content page 1" in result.text
    assert "Test PDF content page 2" in result.text
    assert result.input_hash


def test_extract_pdf_file_not_found():
    """Test PDF extraction with missing file."""
    with pytest.raises(FileNotFoundError):
        extract_pdf(Path("/nonexistent/file.pdf"))


def test_extract_docx(tmp_path):
    """Test DOCX extraction using tmp_path."""
    docx_path = tmp_path / "test.docx"

    doc = DocxDocument()
    doc.add_paragraph("First paragraph of the document")
    doc.add_paragraph("Second paragraph with more content")
    doc.save(str(docx_path))

    result = extract_docx(docx_path)

    assert result.input_kind == "docx"
    assert result.input_name == "test.docx"
    assert "First paragraph" in result.text
    assert "Second paragraph" in result.text
    assert result.input_hash


def test_extract_docx_file_not_found():
    """Test DOCX extraction with missing file."""
    with pytest.raises(FileNotFoundError):
        extract_docx(Path("/nonexistent/file.docx"))


def test_extract_docx_deterministic(tmp_path):
    """Test DOCX extraction is deterministic."""
    docx_path = tmp_path / "test.docx"

    doc = DocxDocument()
    doc.add_paragraph("Deterministic content")
    doc.save(str(docx_path))

    result1 = extract_docx(docx_path)
    result2 = extract_docx(docx_path)

    assert result1.input_hash == result2.input_hash
    assert result1.text == result2.text
