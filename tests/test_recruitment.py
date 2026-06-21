"""Tests for the Recruitment candidate database module."""
from pathlib import Path

import pytest

from core.recruitment.candidate_manager import CandidateManager, CandidateValidationError
from core.recruitment.cv_autofill import autofill_from_text
from core.recruitment.cv_parser import import_cv


def _api_client(tmp_path, monkeypatch):
    monkeypatch.setenv("SAP_DATA_ROOT", str(tmp_path))
    import src.web.dependencies as deps

    monkeypatch.setattr(deps, "DATA_ROOT", tmp_path)
    from src.web.app import app
    from starlette.testclient import TestClient

    return TestClient(app)


class TestCandidateManager:
    def test_schema_and_folders_are_created(self, tmp_path):
        manager = CandidateManager(tmp_path)

        assert manager.db_path == tmp_path / "recruitment" / "candidates.db"
        assert manager.db_path.exists()
        assert manager.cvs_dir.is_dir()

    def test_create_get_update_delete_candidate(self, tmp_path):
        manager = CandidateManager(tmp_path)
        candidate = manager.create_candidate(
            {
                "full_name": "Jane Candidate",
                "email": "jane@example.test",
                "main_role": "SAP IS-U Consultant",
                "work_mode": "Remote",
                "currency": "EUR",
            }
        )

        assert candidate.id
        assert candidate.full_name == "Jane Candidate"
        assert candidate.work_mode == "Remote"
        assert candidate.currency == "EUR"

        updated = manager.update_candidate(
            candidate.id,
            {
                "phone": "+00 600 111 222",
                "skills": "Billing, device management",
                "rate_hour": "55,5",
            },
        )
        assert updated.phone == "+00 600 111 222"
        assert updated.skills == "Billing, device management"
        assert updated.rate_hour == 55.5

        assert manager.get_candidate(candidate.id).email == "jane@example.test"
        assert manager.delete_candidate(candidate.id) is True
        assert manager.get_candidate(candidate.id) is None
        assert manager.delete_candidate(candidate.id) is False

    def test_validation_requires_name_and_contact(self, tmp_path):
        manager = CandidateManager(tmp_path)

        with pytest.raises(CandidateValidationError, match="Full name"):
            manager.create_candidate({"email": "missing-name@example.test"})
        with pytest.raises(CandidateValidationError, match="Email, phone or LinkedIn"):
            manager.create_candidate({"full_name": "No Contact"})
        linkedin_only = manager.create_candidate(
            {"full_name": "Linked Contact", "linkedin": "https://www.linkedin.com/in/linked-contact"}
        )
        assert linkedin_only.linkedin.endswith("/linked-contact")
        with pytest.raises(CandidateValidationError, match="Work mode"):
            manager.create_candidate({"full_name": "Bad Mode", "email": "bad@example.test", "work_mode": "Freelance"})
        with pytest.raises(CandidateValidationError, match="Currency"):
            manager.create_candidate({"full_name": "Bad Currency", "email": "bad@example.test", "currency": "CHF"})
        with pytest.raises(CandidateValidationError, match="numeric"):
            manager.create_candidate({"full_name": "Bad Rate", "email": "bad@example.test", "rate_hour": "abc"})
        with pytest.raises(CandidateValidationError, match="negative"):
            manager.create_candidate({"full_name": "Bad Rate", "email": "bad@example.test", "rate_day": -1})

    def test_search_and_combined_filters(self, tmp_path):
        manager = CandidateManager(tmp_path)
        manager.create_candidate(
            {
                "full_name": "Jane Candidate",
                "email": "jane@example.test",
                "country": "Spain",
                "seniority": "Senior",
                "skills": "ABAP, debugging, invoicing",
                "sap_modules": "SAP IS-U Billing, FI-CA",
                "languages": "English, Spanish",
                "work_mode": "Remote",
                "rate_hour": 60,
                "rate_day": 480,
                "cv_text": "Deep EABL and meter reading background for SAP IS-U projects.",
            }
        )
        manager.create_candidate(
            {
                "full_name": "John Onsite",
                "phone": "+00 7000 000000",
                "country": "United Kingdom",
                "seniority": "Junior",
                "skills": "Testing",
                "sap_modules": "SAP MM",
                "languages": "English",
                "work_mode": "Onsite",
                "rate_hour": 35,
                "rate_day": 300,
            }
        )

        assert [c.full_name for c in manager.list_candidates(search="EABL")] == ["Jane Candidate"]
        assert [c.full_name for c in manager.list_candidates(filters={"cv_text": "IS-U"})] == ["Jane Candidate"]
        assert [c.full_name for c in manager.list_candidates(filters={"cv_text": "meter reading"})] == [
            "Jane Candidate"
        ]
        assert manager.list_candidates(filters={"cv_text": "meter python"}) == []
        filtered = manager.list_candidates(
            filters={
                "cv_text": "EABL",
                "skill_text": "debug",
                "sap_module_text": "billing",
                "language_text": "spanish",
                "country": "spa",
                "seniority": "senior",
                "work_mode": "Remote",
                "max_hourly_rate": "70",
                "max_daily_rate": "500",
            }
        )
        assert [c.full_name for c in filtered] == ["Jane Candidate"]

        assert manager.list_candidates(filters={"max_hourly_rate": "not-a-number"})
        assert manager.list_candidates(filters={"max_hourly_rate": "40"})[0].full_name == "John Onsite"
        assert manager.search_candidates("onsite")[0].full_name == "John Onsite"

    def test_list_order_uses_latest_update_then_id(self, tmp_path):
        manager = CandidateManager(tmp_path)
        first = manager.create_candidate({"full_name": "First Candidate", "email": "first@example.test"})
        second = manager.create_candidate({"full_name": "Second Candidate", "email": "second@example.test"})

        assert [c.id for c in manager.list_candidates()] == [second.id, first.id]


class TestCvParserAndAutofill:
    def test_txt_import_copies_extracts_and_autofills_without_saving(self, tmp_path):
        source = tmp_path / "Jane Candidate CV.txt"
        source.write_text(
            "\n".join(
                [
                    "Jane Candidate",
                    "Senior SAP IS-U Consultant",
                    "jane@example.test | +00 600 111 222",
                    "https://www.linkedin.com/in/jane-candidate",
                    "Over 10 years of experience",
                    "Hourly rate 55 EUR",
                    "Daily rate 450 EUR",
                ]
            ),
            encoding="utf-8",
        )

        imported = import_cv(source, data_root=tmp_path)
        copied = Path(imported["cv_file_path"])
        assert copied.exists()
        assert copied.parent == tmp_path / "recruitment" / "cvs"
        assert "Jane Candidate" in imported["cv_text"]
        assert imported["error"] is None

        data = autofill_from_text(imported["cv_text"])
        assert data["full_name"] == "Jane Candidate"
        assert data["email"] == "jane@example.test"
        assert data["phone"] == "+00 600 111 222"
        assert "linkedin.com/in/jane-candidate" in data["linkedin"]
        assert data["main_role"] == "Senior SAP IS-U Consultant"
        assert data["seniority"] == "Senior"
        assert data["years_experience"] == "over 10 years"
        assert data["rate_hour"] == 55
        assert data["rate_day"] == 450
        assert data["currency"] == "EUR"
        assert "skills" not in data
        assert data["sap_modules"] == "SAP IS-U"

    def test_autofill_handles_columnar_cv_layout_before_name(self):
        text = "\n".join(
            [
                "CONTACT",
                "Phone  +00 700 000 001",
                "Email  candidate.one@example.test",
                "Location  Sample City, Sample Country",
                "LinkedIn  in/candidate-one",
                "SKILLS",
                "SAP UI5 10+ years",
                "SAP HANA 10+ years",
                "BTP 10+ years",
                "CAP 7+ years",
                "CERTIFICATIONS",
                "SAP Certified Application Associate",
                "CANDIDATE ONE",
                "SENIOR SAP FIORI / UI5 / CAP / BTP FULLSTACK DEVELOPER  \u2022",
                "CONTRACTOR",
                "ABOUT ME",
                "I am only interested in remote work and have 10+ years of experience.",
            ]
        )

        data = autofill_from_text(text)

        assert data["full_name"] == "CANDIDATE ONE"
        assert data["email"] == "candidate.one@example.test"
        assert data["phone"] == "+00 700 000 001"
        assert data["linkedin"] == "https://www.linkedin.com/in/candidate-one"
        assert data["city"] == "Sample City"
        assert data["country"] == "Sample Country"
        assert data["main_role"] == "SENIOR SAP FIORI / UI5 / CAP / BTP FULLSTACK DEVELOPER CONTRACTOR"
        assert data["seniority"] == "Senior"
        assert data["years_experience"] == "10+ years"
        assert data["work_mode"] == "Remote"
        assert "SAP UI5" in data["skills"]
        assert "CAP" in data["skills"]
        assert "SAP UI5" in data["sap_modules"]
        assert not data["full_name"].lower().startswith("location")

    def test_autofill_splits_name_role_and_normalizes_spaced_email(self):
        text = "\n".join(
            [
                "Candidate Two  SAP PM / EAM Functional Consultant",
                "candidate.two @example.test",
                "10 Sample Street",
                "Sampletown (France)",
                "+00123456789",
                "SIGNIFICANT PROJECTS (7+ YEARS)",
            ]
        )

        data = autofill_from_text(text)

        assert data["full_name"] == "Candidate Two"
        assert data["main_role"] == "SAP PM / EAM Functional Consultant"
        assert data["email"] == "candidate.two@example.test"
        assert data["phone"] == "+00123456789"
        assert data["city"] == "Sampletown"
        assert data["country"] == "France"
        assert data["seniority"] == "Consultant"
        assert data["years_experience"] == "7+ years"
        assert data["sap_modules"] == "SAP PM / EAM"

    def test_autofill_handles_comma_identity_decimal_years_and_ignores_dates_as_phone(self):
        text = "\n".join(
            [
                "\u00c7agdas Candidate, Senior Abap Developer",
                "Overview",
                "This candidate has 13,5 years experience in SAP world.",
                "Birth Date  : 16.07.1989",
                "Email candidate.three@example.test",
                "SAP Career Summary",
                "\u25aa CDS views",
                "\u25aa SAP BAPI / Function Modules",
                "\u25aa SAP SQL",
                "SAP Work experiences",
                "Role: Developer",
            ]
        )

        data = autofill_from_text(text)

        assert data["full_name"] == "\u00c7agdas Candidate"
        assert data["main_role"] == "Senior Abap Developer"
        assert data["email"] == "candidate.three@example.test"
        assert "phone" not in data
        assert data["seniority"] == "Senior"
        assert data["years_experience"] == "13.5 years"
        assert data["skills"] == "CDS views, SAP BAPI / Function Modules, SAP SQL"
        assert data["sap_modules"] == "SAP BAPI / Function Modules, SAP SQL"

    def test_autofill_does_not_set_currency_from_unrelated_savings_amounts(self):
        text = "\n".join(
            [
                "CANDIDATE FOUR",
                "AI-Augmented Process Transformation Lead",
                "\u20ac 1.1 M+ documented cost savings",
                "14+ years at regulated enterprises",
                "Email: candidate.four@example.test | Phone: +00 30 000 0000",
            ]
        )

        data = autofill_from_text(text)

        assert data["full_name"] == "CANDIDATE FOUR"
        assert data["main_role"] == "AI-Augmented Process Transformation Lead"
        assert data["phone"] == "+00 30 000 0000"
        assert data["years_experience"] == "14+ years"
        assert data["seniority"] == "Lead"
        assert "currency" not in data

    def test_autofill_handles_corrected_operations_cv_without_cross_line_email(self):
        text = "\n".join(
            [
                "Page 1 of 3",
                "Candidate Five",
                "SAP Enterprise Operations & Project Delivery",
                "Experienced across complex enterprise environments.",
                "candidate.five@example.test",
                "+00730034000",
                "Sample City, Romania",
                "WORK EXPERIENCE",
                "Digital Customer Engagement Manager",
                "SAP",
                "SKILLS",
                "SAP S/4 HANA",
                "SAP RISE",
                "SAP Project Systems (PS)",
                "CERTIFICATES",
                "SAP Certified Technology Associate",
            ]
        )

        data = autofill_from_text(text)

        assert data["full_name"] == "Candidate Five"
        assert data["main_role"] == "SAP Enterprise Operations & Project Delivery"
        assert data["email"] == "candidate.five@example.test"
        assert data["phone"] == "+00730034000"
        assert data["city"] == "Sample City"
        assert data["country"] == "Romania"
        assert "Lead" not in data.get("seniority", "")
        assert data["skills"] == "SAP S/4 HANA, SAP RISE, SAP Project Systems (PS)"

    def test_autofill_handles_country_city_then_name_email_same_line(self):
        text = "\n".join(
            [
                "Spain, Seville",
                "Candidate Six        candidate.six@example.test",
                "SAP S/4HANA Technical Consultant | Clean Core | ABAP Cloud Developer",
                "LinkedIn Profile Phone: +00 688990174 Experience Mango: 06/2024 - 06/2026",
                "Technical Skill Set: S/4 HANA, Clean Core, SAP Released APIs, RAP, Eclipse, ABAP, CDS",
                "Built configurable SAP enhancement with RAP-based Fiori app.",
            ]
        )

        data = autofill_from_text(text)

        assert data["full_name"] == "Candidate Six"
        assert data["email"] == "candidate.six@example.test"
        assert data["phone"] == "+00 688990174"
        assert data["country"] == "Spain"
        assert data["city"] == "Seville"
        assert data["main_role"] == "SAP S/4HANA Technical Consultant | Clean Core | ABAP Cloud Developer"
        assert data["seniority"] == "Consultant"
        assert data["skills"] == "S/4 HANA, Clean Core, SAP Released APIs, RAP, Eclipse, ABAP, CDS"

    def test_autofill_handles_spaced_spanish_ocr_text(self):
        text = "\n".join(
            [
                "A l e j a n d r o  M a r t i n  A n t o n a n z a s",
                "A n a l i s t a  p r o g r a m a d o r  s e n i o r  A B A P - I V",
                "T e l e f o n o :  ( + 0 0 )  6 7 8  4 7 0  2 5 9",
                "E m a i l :  a l e x . c a n d i d a t e @ e x a m p l e . t e s t",
                "U b i c a c i o n :  M a d r i d ,  E s p a n a",
                "L i n k e d I n :  w w w . l i n k e d i n . c o m / i n / a l e j a n d r o - c a n d i d a t e",
                "T E C N O L O G I A S",
                "C o r e  y  D B :  A B A P  7 . 5 ,  P O O  /  O O P ,  H A N A",
                "H e r r a m i e n t a s  e  I n t e r f a c e s :  E c l i p s e  A D T ,  S A P  G U I",
            ]
        )

        data = autofill_from_text(text)

        assert data["full_name"] == "Alejandro Martin Antonanzas"
        assert data["main_role"] == "Analista programador senior ABAP-IV"
        assert data["email"] == "alex.candidate@example.test"
        assert data["phone"] == "+00 678 470 259"
        assert data["city"] == "Madrid"
        assert data["country"] == "Spain"
        assert data["linkedin"] == "https://www.linkedin.com/in/alejandro-candidate"
        assert data["seniority"] == "Senior"
        assert "ABAP 7.5" in data["skills"]
        assert "SAP GUI" in data["skills"]

    def test_autofill_handles_remote_timezone_and_split_years(self):
        text = "\n".join(
            [
                "Candidate Seven",
                "SAP ABAP Developer",
                "candidate.seven@example.test +00 634 755 883",
                "PROFILE:",
                "Senior technology professional including +1 2 years dedicated to SAP ecosystems.",
                "Location: Remote (CET / UTC-3).",
                "LANGUAGES: English, Spanish, Portuguese: professional proficiency.",
            ]
        )

        data = autofill_from_text(text)

        assert data["full_name"] == "Candidate Seven"
        assert data["main_role"] == "SAP ABAP Developer"
        assert data["years_experience"] == "12+ years"
        assert data["work_mode"] == "Remote"
        assert data["timezone"] == "CET/UTC-3"
        assert "country" not in data
        assert "city" not in data

    def test_autofill_handles_table_name_linkedin_fallback_and_contact_suffixes(self):
        text = "\n".join(
            [
                "DATOS PERSONALES | DATOS PERSONALES",
                "NOMBRE Y APELLIDOS: | Jose Sample Candidate | NACIONALIDAD: | Sample",
                "Phone:",
                "+",
                "34 641018470",
                "Email sample.candidate@example.comTel",
                "EXPERIENCIA",
                "Analista programador asignado al proyecto CLIENT_A",
            ]
        )

        data = autofill_from_text(text, filename_hint="CV JSample.docx")

        assert data["full_name"] == "Jose Sample Candidate"
        assert data["email"] == "sample.candidate@example.com"
        assert data["phone"] == "+34 641018470"
        assert data["main_role"] == "Analista programador"

    def test_autofill_handles_split_linkedin_languages_and_country_clues(self):
        text = "\n".join(
            [
                "CONTACT",
                "Sevilla, España📍",
                "+34 696 000 000📞",
                "sample.candidate@example.test",
                "linkedin.com/in/sample-candidate🔗",
                "LANGUAGES",
                "Spanish Native",
                "French C1 — Advanced",
                "English B2+ Professional",
                "TECHNOLOGIES",
                "SAP Fiori",
                "David Sample Candidate",
                "Senior SAP Fiori/UI5 Developer | SAP S/4HANA | RAP | CDS Views",
                "ABAP OO",
                "PROFESSIONAL PROFILE",
                "This profile includes over 10 years of experience.",
                "technologies, specialised in developing SAP Fiori and SAPUI5",
            ]
        )

        data = autofill_from_text(text)

        assert data["full_name"] == "David Sample Candidate"
        assert data["linkedin"] == "https://www.linkedin.com/in/sample-candidate"
        assert data["country"] == "Spain"
        assert data["city"] == "Sevilla"
        assert data["main_role"] == "Senior SAP Fiori/UI5 Developer | SAP S/4HANA | RAP | CDS Views ABAP OO"
        assert data["languages"] == "Spanish Native, French C1 — Advanced, English B2+ Professional"

    def test_autofill_handles_multiline_linkedin_and_profile_years_priority(self):
        text = "\n".join(
            [
                "Contact",
                "www.linkedin.com/in/marcelo",
                "sample-99b52718",
                "Marcelo Sample",
                "SAP ABAP Developer",
                "Spain",
                "marcelo.sample@example.test",
                "+34 697346072",
                "PROFILE",
                "SAP ABAP Consultant with 15 years of experience in implementation projects.",
                "Recent project involved over 3 years of maintenance support.",
            ]
        )

        data = autofill_from_text(text)

        assert data["full_name"] == "Marcelo Sample"
        assert data["linkedin"] == "https://www.linkedin.com/in/marcelosample-99b52718"
        assert data["country"] == "Spain"
        assert data["main_role"] == "SAP ABAP Developer"
        assert data["years_experience"] == "15 years"

    def test_autofill_handles_spanish_labeled_name_role_and_local_phone_country(self):
        text = "\n".join(
            [
                "DATOS PERSONALES Y DE CONTACTO",
                "Nombre y Apellidos: Teresa Sample Candidate",
                "Telefono: 630754806",
                "Email: teresa.sample@example.test",
                "Puesto: Freelance ABAP MM/WORKFLOW",
                "ESTUDIOS ACADEMICOS",
                "Ingenieria",
            ]
        )

        data = autofill_from_text(text)

        assert data["full_name"] == "Teresa Sample Candidate"
        assert data["phone"] == "630754806"
        assert data["country"] == "Spain"
        assert data["main_role"] == "Freelance ABAP MM/WORKFLOW"

    def test_autofill_uses_filename_when_summary_has_soft_skill_phrases(self):
        text = "\n".join(
            [
                "PROFESSIONAL SUMMARY",
                "Senior SAP ABAP Consultant with 10+ years of experience in ABAP development.",
                "WORK HISTORY",
                "Senior SAP Consultant, 05/2017 to Current",
                "Performance Improvement.",
                "Telecom-Rome",
                "SAP ABAP Developer, 10/2016 to 12/2016",
                "Data Dictionary.",
                "Drafting of technical analyses",
                "CONTACT",
                "Address: 00010, Sample City Lazio",
                "Phone: +39 3899818719",
                "Email: sample.consultant@example.test",
            ]
        )

        data = autofill_from_text(text, filename_hint="Luca_Sample_Resume.pdf")

        assert data["full_name"] == "Luca Sample"
        assert data["main_role"] == "Senior SAP ABAP Consultant"
        assert data["country"] == "Italy"

    def test_autofill_extracts_role_from_summary_line_with_years_after_name(self):
        text = "\n".join(
            [
                "ARUN SAMPLE",
                "Contact No: +34-631137365 LinkedIn: www.linkedin.com/in/arun-sample Email: sample.abap@example.test",
                "Professional Summary",
                "Senior SAP ABAP Developer with over 11 years of hands-on experience in SAP S/4HANA and ECC environments. Expert in",
                "designing, developing, and optimizing custom SAP solutions.",
                "Brownfield implementations, custom code remediation, and integration projects.",
            ]
        )

        data = autofill_from_text(text)

        assert data["full_name"] == "ARUN SAMPLE"
        assert data["main_role"] == "Senior SAP ABAP Developer"
        assert data["years_experience"] == "over 11 years"

    def test_autofill_handles_surname_name_and_job_function_table_labels(self):
        text = "\n".join(
            [
                "Personal Information",
                "Surname / Name | Silva/Alex | Silva/Alex",
                "Date of Birth | 28/09/1974 | 28/09/1974",
                "Summary",
                "Job or Function | Senior Abap Consultant | Senior Abap Consultant",
                "Profile and Background",
                "SAP Skills | Extensive experience (24 years) in ABAP developments",
            ]
        )

        data = autofill_from_text(text)

        assert data["full_name"] == "Alex Silva"
        assert data["main_role"] == "Senior Abap Consultant"
        assert data["years_experience"] == "24 years"

    def test_autofill_handles_personal_details_before_name_and_skills(self):
        text = "\n".join(
            [
                "PERSONAL DETAILS",
                "Sweden",
                "+46-764340663",
                "sample.abap@example.test",
                "Visa status : Permanent Residence",
                "SKILLS",
                "\uf0b7 SAP RAP(ABAP Restful application programming)",
                "\uf0b7 SAP BTP",
                "\uf0b7 SAP IS-U EDM",
                "Henrik Sample",
                "Over all 19 years of professional experience in SAP ABAP.",
                "EXPERIENCE",
                "Senior SAP Technical Consultant",
            ]
        )

        data = autofill_from_text(text)

        assert data["full_name"] == "Henrik Sample"
        assert data["main_role"] == "Senior SAP Technical Consultant"
        assert data["country"] == "Sweden"
        assert data["skills"] == "SAP RAP(ABAP Restful application programming), SAP BTP, SAP IS-U EDM"

    def test_autofill_prefers_phone_country_over_later_project_locations(self):
        text = "\n".join(
            [
                "Nikos Sample",
                "Senior ABAP/4 Technical Consultant",
                "Address Sample Street, Athens, Greece Mobile +30 6946982207",
                "Email dimitris.sample@example.test",
                "Senior ABAP/4 Developer for CLIENT_A (Spain, Barcelona)",
            ]
        )

        data = autofill_from_text(text)

        assert data["full_name"] == "Nikos Sample"
        assert data["country"] == "Greece"
        assert data["main_role"] == "Senior ABAP/4 Technical Consultant"

    def test_autofill_rejects_company_project_and_module_names(self):
        text = "\n".join(
            [
                "AVANTE IT",
                "Requirement Specification",
                "Material Ledger",
                "Project CLIENT_A",
                "Ricardo Zarate Saenz",
                "Career Summary",
                "Position: SAP Functional FICO-PS Level 3 - S/4 HANA",
                "Over 14 years of experience",
            ]
        )

        data = autofill_from_text(text, filename_hint="CV_Ricardo_Zarate.pdf")

        assert data["full_name"] == "Ricardo Zarate Saenz"
        assert data["main_role"] == "SAP Functional FICO-PS Level 3-S/4 HANA"
        assert data["years_experience"] == "over 14 years"

    def test_autofill_uses_linkedin_and_filename_as_last_resort(self):
        linkedin_text = "LinkedIn profile | https://www.linkedin.com/in/felix-sample-candidate-123abc/"
        linkedin_data = autofill_from_text(linkedin_text, filename_hint="IT CV Felix.docx")

        assert linkedin_data["full_name"] == "Felix Sample Candidate"
        assert linkedin_data["linkedin"].endswith("/felix-sample-candidate-123abc")

        filename_data = autofill_from_text(
            "Business Process Senior Consultant\nEmail sample@example.test\nPhone +00 873585805",
            filename_hint="SAP Resume - Mahesh.pdf",
        )

        assert filename_data["full_name"] == "Mahesh"
        assert filename_data["main_role"] == "Business Process Senior Consultant"

    def test_docx_import_extracts_text(self, tmp_path):
        from docx import Document

        source = tmp_path / "cv.docx"
        document = Document()
        document.add_paragraph("Alex Profile")
        document.add_paragraph("Lead SAP IS-U Architect")
        document.save(source)

        imported = import_cv(source, data_root=tmp_path)

        assert imported["error"] is None
        assert "Alex Profile" in imported["cv_text"]
        assert "Lead SAP IS-U Architect" in imported["cv_text"]

    def test_docx_import_extracts_table_text(self, tmp_path):
        from docx import Document

        source = tmp_path / "table-cv.docx"
        document = Document()
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "NOMBRE Y APELLIDOS:"
        table.cell(0, 1).text = "Jose Sample Candidate"
        document.save(source)

        imported = import_cv(source, data_root=tmp_path)

        assert imported["error"] is None
        assert "NOMBRE Y APELLIDOS" in imported["cv_text"]
        assert "Jose Sample Candidate" in imported["cv_text"]

    def test_pdf_import_extracts_text(self, tmp_path):
        from reportlab.pdfgen import canvas

        source = tmp_path / "cv.pdf"
        pdf = canvas.Canvas(str(source))
        pdf.drawString(72, 720, "PDF Candidate")
        pdf.drawString(72, 700, "SAP IS-U Billing Consultant")
        pdf.save()

        imported = import_cv(source, data_root=tmp_path)

        assert imported["error"] is None
        assert "PDF Candidate" in imported["cv_text"]
        assert "SAP IS-U Billing Consultant" in imported["cv_text"]

    def test_broken_pdf_is_stored_with_friendly_extraction_error(self, tmp_path):
        source = tmp_path / "broken.pdf"
        source.write_bytes(b"not a real pdf")

        imported = import_cv(source, data_root=tmp_path)

        assert Path(imported["cv_file_path"]).exists()
        assert imported["cv_text"] == ""
        assert "text extraction failed" in imported["error"]


class TestRecruitmentApi:
    def test_page_contains_required_controls(self, tmp_path, monkeypatch):
        client = _api_client(tmp_path, monkeypatch)

        resp = client.get("/recruitment")

        assert resp.status_code == 200
        for text in (
            "Recruitment",
            "New Candidate",
            "Edit Candidate",
            "Delete Candidate",
            "Import CV",
            "Open CV",
            "Update CV",
            "CV free text",
            "Searching...",
            "Previous Candidate",
            "Next Candidate",
            "Previous Page",
            "Next Page",
            "Page size",
            "Clear Filters",
            "Refresh",
        ):
            assert text in resp.text
        assert "updateFormCv($event)" in resp.text
        assert '@input.debounce.400ms="loadCandidates()"' in resp.text
        assert "pagedCandidates" in resp.text
        assert "@click.outside" not in resp.text
        assert "externalUrl(candidate.linkedin)" in resp.text
        assert 'accept=".pdf,.doc,.docx,.txt"' in resp.text

    def test_api_crud_filter_and_validation(self, tmp_path, monkeypatch):
        client = _api_client(tmp_path, monkeypatch)

        invalid = client.post("/api/recruitment/candidates", json={"full_name": "No Contact"})
        assert invalid.status_code == 400
        assert "Email, phone or LinkedIn" in invalid.json()["error"]

        created = client.post(
            "/api/recruitment/candidates",
            json={
                "full_name": "Jane Candidate",
                "email": "jane@example.test",
                "country": "Spain",
                "skills": "ABAP debugging",
                "sap_modules": "SAP IS-U Billing",
                "languages": "English, Spanish",
                "work_mode": "Remote",
                "rate_hour": 50,
                "rate_day": 400,
                "cv_text": "Deep EABL meter reading implementation background.",
            },
        )
        assert created.status_code == 200
        candidate = created.json()
        assert candidate["id"]

        listed = client.get(
            "/api/recruitment/candidates",
            params={
                "search": "jane",
                "cv_text": "meter reading",
                "skill_text": "debug",
                "sap_module_text": "billing",
                "language_text": "spanish",
                "country": "Spain",
                "work_mode": "Remote",
                "max_hourly_rate": "55",
                "max_daily_rate": "450",
            },
        )
        assert listed.status_code == 200
        assert listed.json()["count"] == 1

        updated = client.put(
            f"/api/recruitment/candidates/{candidate['id']}",
            json={"phone": "+00 600 111 222", "full_name": "Jane Candidate"},
        )
        assert updated.status_code == 200
        assert updated.json()["phone"] == "+00 600 111 222"
        assert updated.json()["email"] == "jane@example.test"

        deleted = client.delete(f"/api/recruitment/candidates/{candidate['id']}")
        assert deleted.status_code == 200
        assert deleted.json()["deleted"] is True
        assert client.get("/api/recruitment/candidates").json()["count"] == 0

    def test_api_import_cv_prefills_but_does_not_auto_save(self, tmp_path, monkeypatch):
        client = _api_client(tmp_path, monkeypatch)
        content = b"Jane Candidate\nSenior SAP IS-U Consultant\njane@example.test\n"

        resp = client.post(
            "/api/recruitment/cv/import",
            files={"file": ("candidate.txt", content, "text/plain")},
        )

        assert resp.status_code == 200
        data = resp.json()
        assert data["saved"] is False
        assert data["candidate"]["full_name"] == "Jane Candidate"
        assert data["candidate"]["email"] == "jane@example.test"
        assert Path(data["candidate"]["cv_file_path"]).exists()
        assert client.get("/api/recruitment/candidates").json()["count"] == 0

    def test_api_import_rejects_unsupported_files(self, tmp_path, monkeypatch):
        client = _api_client(tmp_path, monkeypatch)

        resp = client.post(
            "/api/recruitment/cv/import",
            files={"file": ("candidate.exe", b"nope", "application/octet-stream")},
        )

        assert resp.status_code == 400
        assert "Unsupported CV file type" in resp.json()["error"]

    def test_open_cv_reports_missing_file_cleanly(self, tmp_path, monkeypatch):
        client = _api_client(tmp_path, monkeypatch)
        created = client.post(
            "/api/recruitment/candidates",
            json={
                "full_name": "Jane Candidate",
                "email": "jane@example.test",
                "cv_file_path": str(tmp_path / "missing.pdf"),
            },
        ).json()

        resp = client.post(f"/api/recruitment/candidates/{created['id']}/open-cv")

        assert resp.status_code == 404
        assert "CV file is missing" in resp.json()["error"]

    def test_ui_metadata_matches_requested_controls(self):
        from ui.recruitment.candidate_form import CURRENCY_OPTIONS, FORM_FIELDS, WORK_MODE_OPTIONS
        from ui.recruitment.candidates_tab import ACTION_BUTTONS, FILTER_FIELDS, TABLE_COLUMNS

        assert "full_name" in FORM_FIELDS
        assert "cv_text" in FORM_FIELDS
        assert WORK_MODE_OPTIONS == ("Remote", "Hybrid", "Onsite", "Any")
        assert CURRENCY_OPTIONS == ("EUR", "USD", "GBP", "Other")
        assert "cv_text" in FILTER_FIELDS
        assert "sap_module_text" in FILTER_FIELDS
        assert "updated_at" in TABLE_COLUMNS
        assert "skills" not in TABLE_COLUMNS
        assert "sap_modules" not in TABLE_COLUMNS
        assert "rate_hour" not in TABLE_COLUMNS
        for action in (
            "New Candidate",
            "Edit Candidate",
            "Delete Candidate",
            "Import CV",
            "Update CV",
            "Open CV",
            "Previous Candidate",
            "Next Candidate",
            "Previous Page",
            "Next Page",
            "Clear Filters",
            "Refresh",
        ):
            assert action in ACTION_BUTTONS
