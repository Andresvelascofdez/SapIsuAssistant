"""
Comprehensive end-to-end tests for ALL functionalities.

Covers:
- Full ingestion pipeline: extract -> synthesize -> store -> approve -> index
- Full chat RAG pipeline: question -> embed -> search -> context -> answer
- OpenAI request structure verification (mocked, no real API key needed)
- OpenAI response handling (mocked responses)
- Incidents saved to DB and read back
- AI retrieves incidents from the system (RAG)
- Kanban full CRUD flow
- Client isolation across all modules
- KB field editing (update_fields)
- Error scenarios: Qdrant down, OpenAI errors, empty inputs
- Token counting and truncation
- Qdrant dimension mismatch detection
- Actionable error messages
- Logging configuration
"""
import json
from pathlib import Path
from unittest.mock import MagicMock, Mock, patch, call

import pytest

from src.assistant.storage.models import KBItem, KBItemType, KBItemStatus, IngestionStatus
from src.assistant.storage.kb_repository import KBItemRepository
from src.assistant.storage.ingestion_repository import IngestionRepository
from src.assistant.ingestion.extractors import extract_text, extract_pdf, extract_docx
from src.assistant.ingestion.synthesis import SynthesisPipeline, SynthesisError, validate_synthesis_output
from src.assistant.chat.chat_service import ChatService, ChatError, ChatResult
from src.assistant.retrieval.qdrant_service import QdrantService
from src.assistant.retrieval.embedding_service import EmbeddingService
from src.kanban.storage.kanban_repository import KanbanRepository, TicketStatus, TicketPriority
from src.shared.client_manager import ClientManager
from src.shared.errors import AppErrors, format_openai_error, format_qdrant_error
from src.shared.tokens import count_tokens, truncate_to_token_limit


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def data_root(tmp_path):
    return tmp_path / "data"


@pytest.fixture
def client_manager(data_root):
    return ClientManager(data_root)


@pytest.fixture
def client_swe(client_manager):
    return client_manager.register_client("SWE", "Stadtwerke Beispiel")


@pytest.fixture
def client_heron(client_manager):
    return client_manager.register_client("HERON", "Heron Energy")


@pytest.fixture
def kb_repo_standard(client_manager):
    std_dir = client_manager.get_standard_dir()
    return KBItemRepository(std_dir / "assistant_kb.sqlite")


@pytest.fixture
def kb_repo_swe(client_manager, client_swe):
    swe_dir = client_manager.get_client_dir("SWE")
    return KBItemRepository(swe_dir / "assistant_kb.sqlite")


@pytest.fixture
def kb_repo_heron(client_manager, client_heron):
    heron_dir = client_manager.get_client_dir("HERON")
    return KBItemRepository(heron_dir / "assistant_kb.sqlite")


@pytest.fixture
def ingestion_repo(client_manager):
    std_dir = client_manager.get_standard_dir()
    return IngestionRepository(std_dir / "assistant_kb.sqlite")


@pytest.fixture
def kanban_swe(client_manager, client_swe):
    swe_dir = client_manager.get_client_dir("SWE")
    return KanbanRepository(swe_dir / "kanban.sqlite")


@pytest.fixture
def kanban_heron(client_manager, client_heron):
    heron_dir = client_manager.get_client_dir("HERON")
    return KanbanRepository(heron_dir / "kanban.sqlite")


def _make_kb_item(repo, title="IDEX Timeout Pattern", item_type=KBItemType.INCIDENT_PATTERN,
                  scope="standard", client_code=None, content="# IDEX Timeout\n\nWhen IDEX times out...",
                  tags=None, sap_objects=None):
    item, is_new = repo.create_or_update(
        client_scope=scope,
        client_code=client_code,
        item_type=item_type,
        title=title,
        content_markdown=content,
        tags=tags or ["IDEX", "timeout"],
        sap_objects=sap_objects or ["EA10", "EL01"],
        signals={"module": "IS-U", "process": "meter-reading"},
        sources={"origin": "manual"},
    )
    return item, is_new


VALID_SYNTHESIS_RESPONSE = json.dumps({
    "kb_items": [
        {
            "type": "INCIDENT_PATTERN",
            "title": "IDEX message timeout in meter reading",
            "content_markdown": "# IDEX Timeout\n\nWhen IDEX messages fail...",
            "tags": ["IDEX", "timeout", "meter-reading"],
            "sap_objects": ["EA10", "EL01", "EDEX_MONITOR"],
            "signals": {"module": "IS-U", "process": "meter-reading", "country": "DE"},
        },
        {
            "type": "RESOLUTION",
            "title": "Restart IDEX processing job",
            "content_markdown": "# Resolution\n\n1. Go to SM37...",
            "tags": ["IDEX", "resolution"],
            "sap_objects": ["SM37", "EA10"],
            "signals": {"module": "IS-U"},
        },
    ]
})


# ===========================================================================
# 1. FULL INGESTION PIPELINE: extract -> synthesize -> store -> approve -> index
# ===========================================================================

class TestFullIngestionPipeline:
    """End-to-end ingestion: text extraction -> OpenAI synthesis -> KB storage -> approve -> Qdrant index."""

    @patch("src.assistant.ingestion.synthesis.OpenAI")
    def test_e2e_text_ingestion(self, mock_openai_class, kb_repo_standard, ingestion_repo):
        """Full pipeline: extract text -> synthesize via OpenAI -> store in KB -> approve."""
        # 1. Extract text
        raw_text = "IDEX timeout during meter reading. Check EA10 and EL01 for stuck messages."
        result = extract_text(raw_text, label="sap_note_001")

        assert result.text == raw_text
        assert result.input_kind == "text"
        assert result.input_name == "sap_note_001"
        assert len(result.input_hash) == 64  # SHA256 hex

        # 2. Record ingestion
        ingestion = ingestion_repo.create(
            client_scope="standard",
            client_code=None,
            input_kind=result.input_kind,
            input_hash=result.input_hash,
            input_name=result.input_name,
            model_used="gpt-5.2",
            reasoning_effort="xhigh",
        )
        assert ingestion.status == IngestionStatus.DRAFT.value

        # 3. Synthesize via OpenAI (mocked)
        mock_client = MagicMock()
        mock_openai_class.return_value = mock_client

        mock_response = Mock()
        mock_response.output_text = VALID_SYNTHESIS_RESPONSE
        mock_client.responses.create.return_value = mock_response

        pipeline = SynthesisPipeline(api_key="test-key")
        synthesis_result = pipeline.synthesize(result.text)

        # Verify OpenAI request was sent correctly
        mock_client.responses.create.assert_called_once()
        call_kwargs = mock_client.responses.create.call_args[1]
        assert call_kwargs["model"] == "gpt-5.2"
        assert call_kwargs["reasoning"]["effort"] == "xhigh"
        assert "json_schema" in str(call_kwargs["text"])
        assert raw_text in call_kwargs["input"]

        # 4. Validate synthesis output
        errors = validate_synthesis_output(synthesis_result)
        assert errors == []
        assert len(synthesis_result["kb_items"]) == 2

        # 5. Store items in KB
        stored_items = []
        for synth_item in synthesis_result["kb_items"]:
            item, is_new = kb_repo_standard.create_or_update(
                client_scope="standard",
                client_code=None,
                item_type=KBItemType(synth_item["type"]),
                title=synth_item["title"],
                content_markdown=synth_item["content_markdown"],
                tags=synth_item["tags"],
                sap_objects=synth_item["sap_objects"],
                signals=synth_item["signals"],
                sources={"ingestion_id": ingestion.ingestion_id},
            )
            assert is_new is True
            assert item.status == KBItemStatus.DRAFT.value
            stored_items.append(item)

        # 6. Update ingestion status
        ingestion_repo.update_status(ingestion.ingestion_id, IngestionStatus.SYNTHESIZED)

        # 7. Approve items
        for item in stored_items:
            kb_repo_standard.update_status(item.kb_id, KBItemStatus.APPROVED)

        ingestion_repo.update_status(ingestion.ingestion_id, IngestionStatus.APPROVED)

        # 8. Verify final state
        for item in stored_items:
            reloaded = kb_repo_standard.get_by_id(item.kb_id)
            assert reloaded.status == KBItemStatus.APPROVED.value

        final_ingestion = ingestion_repo.get_by_id(ingestion.ingestion_id)
        assert final_ingestion.status == IngestionStatus.APPROVED.value

    def test_e2e_pdf_extraction(self, tmp_path):
        """PDF extraction produces valid result."""
        from reportlab.pdfgen import canvas

        pdf_path = tmp_path / "test.pdf"
        c = canvas.Canvas(str(pdf_path))
        c.drawString(100, 750, "SAP IS-U IDEX configuration guide")
        c.drawString(100, 730, "Transaction EA10 is used for meter reading")
        c.save()

        result = extract_pdf(pdf_path)
        assert "EA10" in result.text
        assert result.input_kind == "pdf"
        assert result.input_name == "test.pdf"

    def test_e2e_docx_extraction(self, tmp_path):
        """DOCX extraction produces valid result."""
        from docx import Document

        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("SAP IS-U meter reading troubleshooting")
        doc.add_paragraph("Check transaction EL01 for stuck IDEX messages")
        doc.save(str(docx_path))

        result = extract_docx(docx_path)
        assert "EL01" in result.text
        assert result.input_kind == "docx"


# ===========================================================================
# 2. FULL CHAT RAG PIPELINE: question -> embed -> search -> context -> answer
# ===========================================================================

class TestFullChatRAGPipeline:
    """End-to-end chat RAG: embed question -> Qdrant search -> fetch KB -> context -> OpenAI answer."""

    @patch("src.assistant.chat.chat_service.OpenAI")
    @patch("src.assistant.retrieval.embedding_service.OpenAI")
    @patch("src.assistant.retrieval.qdrant_service.QdrantClient")
    def test_e2e_chat_answer(self, mock_qdrant_class, mock_embed_openai, mock_chat_openai,
                             kb_repo_standard):
        """Full RAG pipeline from question to answer."""
        # Setup: create and approve a KB item
        item, _ = _make_kb_item(kb_repo_standard)
        kb_repo_standard.update_status(item.kb_id, KBItemStatus.APPROVED)

        # Mock embedding service
        mock_embed_client = MagicMock()
        mock_embed_openai.return_value = mock_embed_client
        mock_embed_response = Mock()
        mock_embed_response.data = [Mock(embedding=[0.1] * 3072)]
        mock_embed_client.embeddings.create.return_value = mock_embed_response

        # Mock Qdrant search
        mock_qdrant = MagicMock()
        mock_qdrant_class.return_value = mock_qdrant
        mock_qdrant.collection_exists.return_value = True
        mock_info = Mock()
        mock_info.config.params.vectors.size = 3072
        mock_qdrant.get_collection.return_value = mock_info
        mock_hit = Mock()
        mock_hit.payload = {"kb_id": item.kb_id}
        mock_hit.score = 0.95
        mock_qdrant.search.return_value = [mock_hit]

        # Mock OpenAI chat response
        mock_chat_client = MagicMock()
        mock_chat_openai.return_value = mock_chat_client
        mock_chat_response = Mock()
        mock_chat_response.output_text = "To fix the IDEX timeout, check EA10 for stuck messages."
        mock_chat_client.responses.create.return_value = mock_chat_response

        # Execute
        embed_svc = EmbeddingService(api_key="test-key")
        qdrant_svc = QdrantService("http://localhost:6333")
        chat_svc = ChatService(embed_svc, qdrant_svc, api_key="test-key")

        result = chat_svc.answer(
            question="How do I fix IDEX timeout?",
            kb_repo=kb_repo_standard,
            client_scope="standard",
            client_code=None,
            include_standard=True,
        )

        # Verify result
        assert isinstance(result, ChatResult)
        assert "IDEX" in result.answer
        assert len(result.sources) == 1
        assert result.sources[0].kb_id == item.kb_id

        # Verify OpenAI embedding request was sent
        mock_embed_client.embeddings.create.assert_called_once()
        embed_call = mock_embed_client.embeddings.create.call_args[1]
        assert embed_call["model"] == "text-embedding-3-large"
        assert "IDEX timeout" in embed_call["input"]

        # Verify OpenAI chat request was sent with correct structure
        mock_chat_client.responses.create.assert_called_once()
        chat_call = mock_chat_client.responses.create.call_args[1]
        assert chat_call["model"] == "gpt-5.2"
        assert "SAP IS-U technical assistant" in chat_call["instructions"]
        assert "IDEX timeout" in chat_call["input"]
        assert "reasoning" in chat_call

    @patch("src.assistant.chat.chat_service.OpenAI")
    @patch("src.assistant.retrieval.embedding_service.OpenAI")
    @patch("src.assistant.retrieval.qdrant_service.QdrantClient")
    def test_chat_with_multiple_sources(self, mock_qdrant_class, mock_embed_openai,
                                         mock_chat_openai, kb_repo_standard):
        """Chat with multiple KB items as context."""
        item1, _ = _make_kb_item(kb_repo_standard, title="IDEX Timeout", content="Timeout issue...")
        item2, _ = _make_kb_item(kb_repo_standard, title="IDEX Resolution",
                                  item_type=KBItemType.RESOLUTION, content="Fix: restart job in SM37")
        kb_repo_standard.update_status(item1.kb_id, KBItemStatus.APPROVED)
        kb_repo_standard.update_status(item2.kb_id, KBItemStatus.APPROVED)

        # Mock services
        mock_embed_client = MagicMock()
        mock_embed_openai.return_value = mock_embed_client
        mock_embed_client.embeddings.create.return_value = Mock(data=[Mock(embedding=[0.1] * 3072)])

        mock_qdrant = MagicMock()
        mock_qdrant_class.return_value = mock_qdrant
        mock_qdrant.collection_exists.return_value = True
        mock_info = Mock()
        mock_info.config.params.vectors.size = 3072
        mock_qdrant.get_collection.return_value = mock_info

        hit1 = Mock(payload={"kb_id": item1.kb_id}, score=0.95)
        hit2 = Mock(payload={"kb_id": item2.kb_id}, score=0.90)
        mock_qdrant.search.return_value = [hit1, hit2]

        mock_chat_client = MagicMock()
        mock_chat_openai.return_value = mock_chat_client
        mock_chat_client.responses.create.return_value = Mock(
            output_text="Step 1: Check EA10. Step 2: Restart SM37 job."
        )

        embed_svc = EmbeddingService(api_key="test-key")
        qdrant_svc = QdrantService()
        chat_svc = ChatService(embed_svc, qdrant_svc, api_key="test-key")

        result = chat_svc.answer(
            question="How to resolve IDEX timeout?",
            kb_repo=kb_repo_standard,
            client_scope="standard",
            client_code=None,
        )

        assert len(result.sources) == 2
        assert "Step 1" in result.answer

    @patch("src.assistant.chat.chat_service.OpenAI")
    @patch("src.assistant.retrieval.embedding_service.OpenAI")
    @patch("src.assistant.retrieval.qdrant_service.QdrantClient")
    def test_chat_no_results_returns_missing_context(self, mock_qdrant_class, mock_embed_openai,
                                                      mock_chat_openai, kb_repo_standard):
        """Chat with no search results still calls OpenAI with 'no relevant items' context."""
        mock_embed_client = MagicMock()
        mock_embed_openai.return_value = mock_embed_client
        mock_embed_client.embeddings.create.return_value = Mock(data=[Mock(embedding=[0.1] * 3072)])

        mock_qdrant = MagicMock()
        mock_qdrant_class.return_value = mock_qdrant
        mock_qdrant.collection_exists.return_value = False
        mock_qdrant.search.return_value = []

        mock_chat_client = MagicMock()
        mock_chat_openai.return_value = mock_chat_client
        mock_chat_client.responses.create.return_value = Mock(
            output_text="I don't have enough context to answer."
        )

        embed_svc = EmbeddingService(api_key="test-key")
        qdrant_svc = QdrantService()
        chat_svc = ChatService(embed_svc, qdrant_svc, api_key="test-key")

        result = chat_svc.answer(
            question="Unknown topic",
            kb_repo=kb_repo_standard,
            client_scope="standard",
            client_code=None,
        )

        # Verify context pack says no items
        chat_call = mock_chat_client.responses.create.call_args[1]
        assert "No relevant knowledge items found" in chat_call["input"]


# ===========================================================================
# 3. OPENAI REQUEST/RESPONSE VERIFICATION
# ===========================================================================

class TestOpenAIRequestResponse:
    """Verify OpenAI requests are correctly structured and responses handled."""

    @patch("src.assistant.ingestion.synthesis.OpenAI")
    def test_synthesis_request_structure(self, mock_openai_class):
        """Verify synthesis sends correct request to OpenAI Responses API."""
        mock_client = MagicMock()
        mock_openai_class.return_value = mock_client
        mock_client.responses.create.return_value = Mock(output_text=VALID_SYNTHESIS_RESPONSE)

        pipeline = SynthesisPipeline(api_key="sk-test-key-123")

        pipeline.synthesize("Sample SAP text about IDEX")

        # Verify request structure
        mock_client.responses.create.assert_called_once()
        kwargs = mock_client.responses.create.call_args[1]

        assert kwargs["model"] == "gpt-5.2"
        assert kwargs["reasoning"] == {"effort": "xhigh"}
        assert "instructions" in kwargs
        assert "SAP IS-U knowledge engineer" in kwargs["instructions"]
        assert "IDEX" in kwargs["input"]
        assert kwargs["text"]["format"]["type"] == "json_schema"
        assert kwargs["text"]["format"]["name"] == "kb_synthesis"

    @patch("src.assistant.ingestion.synthesis.OpenAI")
    def test_synthesis_handles_malformed_json_response(self, mock_openai_class):
        """Synthesis handles JSON decode errors gracefully."""
        mock_client = MagicMock()
        mock_openai_class.return_value = mock_client
        mock_client.responses.create.return_value = Mock(output_text="not valid json {{{")

        pipeline = SynthesisPipeline(api_key="test-key")

        with pytest.raises(SynthesisError, match="failed after"):
            pipeline.synthesize("test text")

    @patch("src.assistant.ingestion.synthesis.OpenAI")
    def test_synthesis_retries_on_invalid_schema(self, mock_openai_class):
        """Synthesis retries once when schema validation fails."""
        mock_client = MagicMock()
        mock_openai_class.return_value = mock_client

        # First call: invalid (missing kb_items), second call: valid
        mock_client.responses.create.side_effect = [
            Mock(output_text='{"wrong_key": []}'),
            Mock(output_text=VALID_SYNTHESIS_RESPONSE),
        ]

        pipeline = SynthesisPipeline(api_key="test-key")
        result = pipeline.synthesize("test text")

        assert mock_client.responses.create.call_count == 2
        assert "kb_items" in result

    @patch("src.assistant.retrieval.embedding_service.OpenAI")
    def test_embedding_request_structure(self, mock_openai_class):
        """Verify embedding request uses text-embedding-3-large."""
        mock_client = MagicMock()
        mock_openai_class.return_value = mock_client
        mock_client.embeddings.create.return_value = Mock(data=[Mock(embedding=[0.5] * 3072)])

        svc = EmbeddingService(api_key="test-key")
        result = svc.embed("Test text for embedding")

        mock_client.embeddings.create.assert_called_once_with(
            model="text-embedding-3-large",
            input="Test text for embedding",
        )
        assert len(result) == 3072

    @patch("src.assistant.retrieval.embedding_service.OpenAI")
    def test_batch_embedding_request(self, mock_openai_class):
        """Verify batch embedding sends multiple texts in one request."""
        mock_client = MagicMock()
        mock_openai_class.return_value = mock_client
        mock_client.embeddings.create.return_value = Mock(
            data=[Mock(embedding=[0.1] * 3072), Mock(embedding=[0.2] * 3072)]
        )

        svc = EmbeddingService(api_key="test-key")
        results = svc.embed_batch(["text1", "text2"])

        mock_client.embeddings.create.assert_called_once_with(
            model="text-embedding-3-large",
            input=["text1", "text2"],
        )
        assert len(results) == 2

    @patch("src.assistant.chat.chat_service.OpenAI")
    @patch("src.assistant.retrieval.embedding_service.OpenAI")
    @patch("src.assistant.retrieval.qdrant_service.QdrantClient")
    def test_chat_response_with_reasoning_effort(self, mock_qdrant_class, mock_embed_openai,
                                                   mock_chat_openai, kb_repo_standard):
        """Verify chat sends correct reasoning effort parameter."""
        mock_embed_client = MagicMock()
        mock_embed_openai.return_value = mock_embed_client
        mock_embed_client.embeddings.create.return_value = Mock(data=[Mock(embedding=[0.1] * 3072)])

        mock_qdrant = MagicMock()
        mock_qdrant_class.return_value = mock_qdrant
        mock_qdrant.collection_exists.return_value = False

        mock_chat_client = MagicMock()
        mock_chat_openai.return_value = mock_chat_client
        mock_chat_client.responses.create.return_value = Mock(output_text="Answer")

        embed_svc = EmbeddingService(api_key="test-key")
        qdrant_svc = QdrantService()
        chat_svc = ChatService(embed_svc, qdrant_svc, api_key="test-key")

        chat_svc.answer(
            question="test",
            kb_repo=kb_repo_standard,
            client_scope="standard",
            client_code=None,
            reasoning_effort="xhigh",
        )

        chat_call = mock_chat_client.responses.create.call_args[1]
        assert chat_call["reasoning"] == {"effort": "xhigh"}


# ===========================================================================
# 4. INCIDENTS SAVED TO DB AND READ BACK
# ===========================================================================

class TestKBItemPersistence:
    """Verify KB items (incidents) are correctly saved and read from SQLite."""

    def test_create_and_read_incident(self, kb_repo_standard):
        item, is_new = _make_kb_item(kb_repo_standard)
        assert is_new is True

        reloaded = kb_repo_standard.get_by_id(item.kb_id)
        assert reloaded is not None
        assert reloaded.title == "IDEX Timeout Pattern"
        assert reloaded.type == KBItemType.INCIDENT_PATTERN.value
        assert reloaded.content_markdown == "# IDEX Timeout\n\nWhen IDEX times out..."
        assert json.loads(reloaded.tags_json) == ["IDEX", "timeout"]
        assert json.loads(reloaded.sap_objects_json) == ["EA10", "EL01"]
        assert reloaded.version == 1
        assert reloaded.status == KBItemStatus.DRAFT.value

    def test_dedupe_same_content(self, kb_repo_standard):
        """Same type + title + content = no duplicate."""
        item1, new1 = _make_kb_item(kb_repo_standard)
        item2, new2 = _make_kb_item(kb_repo_standard)

        assert new1 is True
        assert new2 is False
        assert item1.kb_id == item2.kb_id

    def test_version_different_content(self, kb_repo_standard):
        """Same type + title + different content = new version."""
        item1, new1 = _make_kb_item(kb_repo_standard, content="Version 1 content")
        item2, new2 = _make_kb_item(kb_repo_standard, content="Version 2 content")

        assert new1 is True
        assert new2 is False
        assert item2.version == 2
        assert item2.kb_id == item1.kb_id

    def test_list_by_scope_standard(self, kb_repo_standard):
        _make_kb_item(kb_repo_standard, title="Item 1")
        _make_kb_item(kb_repo_standard, title="Item 2")

        items = kb_repo_standard.list_by_scope("standard")
        assert len(items) == 2

    def test_list_by_scope_with_status_filter(self, kb_repo_standard):
        item, _ = _make_kb_item(kb_repo_standard)
        kb_repo_standard.update_status(item.kb_id, KBItemStatus.APPROVED)
        _make_kb_item(kb_repo_standard, title="Draft item")

        approved = kb_repo_standard.list_by_scope("standard", status=KBItemStatus.APPROVED)
        drafts = kb_repo_standard.list_by_scope("standard", status=KBItemStatus.DRAFT)

        assert len(approved) == 1
        assert len(drafts) == 1

    def test_update_fields_persists_edits(self, kb_repo_standard):
        """update_fields persists title, tags, sap_objects, content changes."""
        item, _ = _make_kb_item(kb_repo_standard)

        updated = kb_repo_standard.update_fields(
            item.kb_id,
            title="Updated Title",
            content_markdown="# Updated\n\nNew content",
            tags=["new-tag1", "new-tag2"],
            sap_objects=["SM37"],
        )

        assert updated.title == "Updated Title"
        assert updated.content_markdown == "# Updated\n\nNew content"
        assert json.loads(updated.tags_json) == ["new-tag1", "new-tag2"]
        assert json.loads(updated.sap_objects_json) == ["SM37"]
        assert updated.content_hash != item.content_hash  # Hash recomputed

    def test_update_fields_partial(self, kb_repo_standard):
        """update_fields with only some fields leaves others unchanged."""
        item, _ = _make_kb_item(kb_repo_standard)

        updated = kb_repo_standard.update_fields(item.kb_id, title="New Title Only")

        assert updated.title == "New Title Only"
        assert updated.content_markdown == item.content_markdown  # Unchanged
        assert updated.tags_json == item.tags_json  # Unchanged

    def test_update_fields_nonexistent(self, kb_repo_standard):
        result = kb_repo_standard.update_fields("nonexistent-id", title="X")
        assert result is None

    def test_all_kb_item_types_stored(self, kb_repo_standard):
        """All 8 KBItemType values can be stored."""
        for item_type in KBItemType:
            item, is_new = _make_kb_item(
                kb_repo_standard,
                title=f"Test {item_type.value}",
                item_type=item_type,
            )
            assert is_new is True
            assert item.type == item_type.value

        items = kb_repo_standard.list_by_scope("standard")
        assert len(items) == 8


# ===========================================================================
# 5. AI RETRIEVES INCIDENTS FROM THE SYSTEM (RAG)
# ===========================================================================

class TestAIRetrievesIncidents:
    """Verify the full RAG flow: AI can retrieve stored incidents."""

    def test_context_pack_includes_incident_data(self, kb_repo_standard):
        """Context pack contains incident title, type, tags, content."""
        item, _ = _make_kb_item(kb_repo_standard)

        source_items = [(item, 0.95)]
        context = ChatService._build_context_pack(source_items)

        assert "IDEX Timeout Pattern" in context
        assert "INCIDENT_PATTERN" in context
        assert "IDEX" in context
        assert "EA10" in context
        assert "Score: 0.950" in context

    def test_context_pack_multiple_items_ordered(self, kb_repo_standard):
        """Context pack numbers items sequentially."""
        item1, _ = _make_kb_item(kb_repo_standard, title="First Item")
        item2, _ = _make_kb_item(kb_repo_standard, title="Second Item",
                                  item_type=KBItemType.RESOLUTION)

        source_items = [(item1, 0.95), (item2, 0.85)]
        context = ChatService._build_context_pack(source_items)

        assert "[1]" in context
        assert "[2]" in context
        assert context.index("First Item") < context.index("Second Item")

    def test_context_pack_token_truncation(self, kb_repo_standard):
        """Context pack respects token budget."""
        long_content = "Word " * 50000  # Very long content
        item, _ = _make_kb_item(kb_repo_standard, content=long_content)

        source_items = [(item, 0.95)]
        context = ChatService._build_context_pack(source_items, max_tokens=500)

        tokens = count_tokens(context)
        assert tokens <= 600  # Some slack for truncation granularity


# ===========================================================================
# 6. KANBAN FULL CRUD FLOW
# ===========================================================================

class TestKanbanFullCRUD:
    """End-to-end Kanban ticket lifecycle."""

    def test_create_ticket(self, kanban_swe):
        ticket = kanban_swe.create_ticket(
            title="Fix IDEX timeout",
            priority=TicketPriority.HIGH,
            notes="Critical issue affecting meter reading",
        )

        assert ticket.title == "Fix IDEX timeout"
        assert ticket.status == TicketStatus.EN_PROGRESO
        assert ticket.priority == TicketPriority.HIGH
        assert ticket.notes == "Critical issue affecting meter reading"

    def test_full_ticket_lifecycle(self, kanban_swe):
        """EN_PROGRESO -> ANALIZADO -> TESTING -> ANALIZADO -> CERRADO."""
        ticket = kanban_swe.create_ticket(title="IDEX Fix", priority=TicketPriority.HIGH)
        assert ticket.status == TicketStatus.EN_PROGRESO

        ticket = kanban_swe.update_status(ticket.id, TicketStatus.ANALIZADO)
        assert ticket.status == TicketStatus.ANALIZADO

        ticket = kanban_swe.update_status(ticket.id, TicketStatus.TESTING)
        assert ticket.status == TicketStatus.TESTING

        ticket = kanban_swe.update_status(ticket.id, TicketStatus.ANALIZADO)
        assert ticket.status == TicketStatus.ANALIZADO

        ticket = kanban_swe.update_status(ticket.id, TicketStatus.CERRADO)
        assert ticket.status == TicketStatus.CERRADO
        assert ticket.closed_at is not None

        # Verify history
        history = kanban_swe.get_history(ticket.id)
        assert len(history) == 5  # create + 4 transitions
        assert history[0].from_status is None
        assert history[0].to_status == TicketStatus.EN_PROGRESO
        assert history[-1].to_status == TicketStatus.CERRADO

    def test_update_ticket_fields(self, kanban_swe):
        ticket = kanban_swe.create_ticket(title="Original Title")

        updated = kanban_swe.update_ticket(
            ticket.id,
            title="Updated Title",
            priority=TicketPriority.CRITICAL,
            notes="New notes",
        )

        assert updated.title == "Updated Title"
        assert updated.priority == TicketPriority.CRITICAL
        assert updated.notes == "New notes"

    def test_list_and_filter_tickets(self, kanban_swe):
        kanban_swe.create_ticket(title="T1")
        t2 = kanban_swe.create_ticket(title="T2")
        kanban_swe.update_status(t2.id, TicketStatus.ANALIZADO)

        all_tickets = kanban_swe.list_tickets()
        open_tickets = kanban_swe.list_tickets(status=TicketStatus.EN_PROGRESO)
        ip_tickets = kanban_swe.list_tickets(status=TicketStatus.ANALIZADO)

        assert len(all_tickets) == 2
        assert len(open_tickets) == 1
        assert len(ip_tickets) == 1


# ===========================================================================
# 7. CLIENT ISOLATION ACROSS ALL MODULES
# ===========================================================================

class TestClientIsolation:
    """Verify strict client isolation: data never crosses clients."""

    def test_kb_items_isolated_between_clients(self, kb_repo_swe, kb_repo_heron):
        """KB items in SWE are not visible from HERON."""
        _make_kb_item(kb_repo_swe, title="SWE Incident", scope="client", client_code="SWE")
        _make_kb_item(kb_repo_heron, title="HERON Incident", scope="client", client_code="HERON")

        swe_items = kb_repo_swe.list_by_scope("client", client_code="SWE")
        heron_items = kb_repo_heron.list_by_scope("client", client_code="HERON")

        assert len(swe_items) == 1
        assert swe_items[0].title == "SWE Incident"
        assert len(heron_items) == 1
        assert heron_items[0].title == "HERON Incident"

    def test_kanban_isolated_between_clients(self, kanban_swe, kanban_heron):
        """Kanban tickets in SWE are not visible from HERON."""
        kanban_swe.create_ticket(title="SWE Ticket")
        kanban_heron.create_ticket(title="HERON Ticket")

        swe_tickets = kanban_swe.list_tickets()
        heron_tickets = kanban_heron.list_tickets()

        assert len(swe_tickets) == 1
        assert swe_tickets[0].title == "SWE Ticket"
        assert len(heron_tickets) == 1
        assert heron_tickets[0].title == "HERON Ticket"

    def test_client_folders_isolated(self, client_manager, client_swe, client_heron):
        """Each client has its own directory structure."""
        swe_dir = client_manager.get_client_dir("SWE")
        heron_dir = client_manager.get_client_dir("HERON")

        assert swe_dir != heron_dir
        assert swe_dir.exists()
        assert heron_dir.exists()
        assert (swe_dir / "assistant_kb.sqlite").exists()
        assert (heron_dir / "assistant_kb.sqlite").exists()
        assert (swe_dir / "kanban.sqlite").exists()
        assert (heron_dir / "kanban.sqlite").exists()

    def test_standard_scope_separate_from_clients(self, client_manager, kb_repo_standard, kb_repo_swe):
        """Standard KB items are separate from client items."""
        _make_kb_item(kb_repo_standard, title="Standard Item", scope="standard")
        _make_kb_item(kb_repo_swe, title="Client Item", scope="client", client_code="SWE")

        std_items = kb_repo_standard.list_by_scope("standard")
        swe_items = kb_repo_swe.list_by_scope("client", client_code="SWE")

        assert len(std_items) == 1
        assert std_items[0].title == "Standard Item"
        assert len(swe_items) == 1
        assert swe_items[0].title == "Client Item"

    @patch("src.assistant.retrieval.qdrant_service.QdrantClient")
    def test_qdrant_collections_per_client(self, mock_qdrant_class):
        """Qdrant uses separate collections per client."""
        mock_qdrant = MagicMock()
        mock_qdrant_class.return_value = mock_qdrant
        mock_qdrant.collection_exists.return_value = False

        svc = QdrantService()

        svc.ensure_collection_exists("standard", None)
        svc.ensure_collection_exists("client", "SWE")
        svc.ensure_collection_exists("client", "HERON")

        calls = mock_qdrant.create_collection.call_args_list
        collection_names = [c[1]["collection_name"] for c in calls]

        assert "kb_standard" in collection_names
        assert "kb_SWE" in collection_names
        assert "kb_HERON" in collection_names


# ===========================================================================
# 8. ERROR SCENARIOS
# ===========================================================================

class TestErrorScenarios:
    """Test error handling: Qdrant down, OpenAI errors, empty inputs."""

    def test_empty_text_extraction_raises(self):
        with pytest.raises(ValueError, match="empty"):
            extract_text("")

    def test_empty_whitespace_extraction_raises(self):
        with pytest.raises(ValueError, match="empty"):
            extract_text("   \n\t  ")

    def test_nonexistent_pdf_raises(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            extract_pdf(tmp_path / "nonexistent.pdf")

    def test_nonexistent_docx_raises(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            extract_docx(tmp_path / "nonexistent.docx")

    @patch("src.assistant.chat.chat_service.OpenAI")
    @patch("src.assistant.retrieval.embedding_service.OpenAI")
    def test_chat_error_on_embedding_failure(self, mock_embed_openai, mock_chat_openai,
                                              kb_repo_standard):
        """Chat raises ChatError with actionable message when embedding fails."""
        mock_embed_client = MagicMock()
        mock_embed_openai.return_value = mock_embed_client
        from openai import AuthenticationError
        mock_embed_client.embeddings.create.side_effect = AuthenticationError(
            message="Invalid API key",
            response=Mock(status_code=401),
            body=None,
        )

        mock_qdrant_class_patched = patch("src.assistant.retrieval.qdrant_service.QdrantClient")
        mock_qdrant_class = mock_qdrant_class_patched.start()
        mock_qdrant_class.return_value = MagicMock()

        embed_svc = EmbeddingService(api_key="bad-key")
        qdrant_svc = QdrantService()
        chat_svc = ChatService(embed_svc, qdrant_svc, api_key="bad-key")

        with pytest.raises(ChatError) as exc_info:
            chat_svc.answer("test", kb_repo_standard, "standard", None)

        assert "API key" in str(exc_info.value)
        mock_qdrant_class_patched.stop()

    @patch("src.assistant.chat.chat_service.OpenAI")
    @patch("src.assistant.retrieval.embedding_service.OpenAI")
    @patch("src.assistant.retrieval.qdrant_service.QdrantClient")
    def test_chat_error_on_qdrant_failure(self, mock_qdrant_class, mock_embed_openai,
                                           mock_chat_openai, kb_repo_standard):
        """Chat raises ChatError with actionable message when Qdrant is down."""
        mock_embed_client = MagicMock()
        mock_embed_openai.return_value = mock_embed_client
        mock_embed_client.embeddings.create.return_value = Mock(data=[Mock(embedding=[0.1] * 3072)])

        mock_qdrant = MagicMock()
        mock_qdrant_class.return_value = mock_qdrant
        mock_qdrant.collection_exists.return_value = True
        mock_info = Mock()
        mock_info.config.params.vectors.size = 3072
        mock_qdrant.get_collection.return_value = mock_info
        mock_qdrant.search.side_effect = Exception("Connection refused to localhost:6333")

        embed_svc = EmbeddingService(api_key="test-key")
        qdrant_svc = QdrantService()
        chat_svc = ChatService(embed_svc, qdrant_svc, api_key="test-key")

        with pytest.raises(ChatError) as exc_info:
            chat_svc.answer("test", kb_repo_standard, "standard", None)

        assert "Qdrant" in str(exc_info.value)

    def test_format_openai_auth_error(self):
        from openai import AuthenticationError
        err = AuthenticationError(message="bad key", response=Mock(status_code=401), body=None)
        msg = format_openai_error(err)
        assert msg == AppErrors.OPENAI_AUTH_FAILED

    def test_format_openai_rate_limit_error(self):
        from openai import RateLimitError
        err = RateLimitError(message="rate limit", response=Mock(status_code=429), body=None)
        msg = format_openai_error(err)
        assert msg == AppErrors.OPENAI_RATE_LIMIT

    def test_format_qdrant_connection_error(self):
        err = ConnectionError("Connection refused")
        msg = format_qdrant_error(err)
        assert msg == AppErrors.QDRANT_UNREACHABLE

    def test_format_qdrant_generic_error(self):
        msg = format_qdrant_error(ValueError("some qdrant issue"))
        assert "Qdrant error" in msg

    @patch("src.assistant.retrieval.qdrant_service.QdrantClient")
    def test_qdrant_dimension_mismatch_raises(self, mock_qdrant_class):
        """Qdrant raises when collection has wrong vector dimensions."""
        mock_qdrant = MagicMock()
        mock_qdrant_class.return_value = mock_qdrant
        mock_qdrant.collection_exists.return_value = True

        mock_info = Mock()
        mock_info.config.params.vectors.size = 1536  # Wrong size
        mock_qdrant.get_collection.return_value = mock_info

        svc = QdrantService()

        with pytest.raises(ValueError, match="vector size mismatch"):
            svc.ensure_collection_exists("standard", None)

    def test_duplicate_client_registration_raises(self, client_manager):
        client_manager.register_client("DUP", "Duplicate")
        with pytest.raises(ValueError, match="already exists"):
            client_manager.register_client("DUP", "Duplicate Again")

    def test_empty_client_code_raises(self, client_manager):
        with pytest.raises(ValueError, match="empty"):
            client_manager.register_client("", "No Code")

    def test_empty_client_name_raises(self, client_manager):
        with pytest.raises(ValueError, match="empty"):
            client_manager.register_client("TEST", "")


# ===========================================================================
# 9. TOKEN COUNTING AND TRUNCATION
# ===========================================================================

class TestTokens:
    """Test tiktoken-based token counting and truncation."""

    def test_count_tokens_basic(self):
        tokens = count_tokens("Hello world")
        assert tokens > 0
        assert isinstance(tokens, int)

    def test_count_tokens_empty(self):
        assert count_tokens("") == 0

    def test_truncate_within_limit(self):
        text = "Short text"
        result = truncate_to_token_limit(text, 1000)
        assert result == text

    def test_truncate_over_limit(self):
        text = "word " * 10000  # Very long text
        result = truncate_to_token_limit(text, 100)
        result_tokens = count_tokens(result)
        assert result_tokens <= 100

    def test_truncate_preserves_start(self):
        text = "START " + "filler " * 10000 + " END"
        result = truncate_to_token_limit(text, 50)
        assert result.startswith("START")


# ===========================================================================
# 10. LOGGING CONFIGURATION
# ===========================================================================

class TestLogging:
    """Test structured logging setup."""

    def test_configure_logging_runs(self):
        from src.shared.logging_config import configure_logging
        configure_logging()  # Should not raise

    def test_configure_logging_sets_level(self):
        import logging
        from src.shared.logging_config import configure_logging
        configure_logging(level=logging.DEBUG)
        assert logging.getLogger().level == logging.DEBUG


# ===========================================================================
# 11. ENTRY POINT
# ===========================================================================

class TestEntryPoint:
    """Test __main__.py entry point."""

    def test_main_module_importable(self):
        import src.__main__
        assert hasattr(src.__main__, "main")
